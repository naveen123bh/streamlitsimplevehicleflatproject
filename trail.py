import os
import re
import json
import hashlib
from pathlib import Path
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from ddgs import DDGS
from dotenv import load_dotenv
from docx import Document
from ollama import chat
from playwright.sync_api import sync_playwright


# ============================================================
# CONFIGURATION
# ============================================================

load_dotenv()

MODEL = os.getenv("MODEL_NAME", "qwen3:1.7b")

BASE_DIR = Path(__file__).resolve().parent

PROFILE_FILE = BASE_DIR / "profile.json"
RESUME_DIR = BASE_DIR / "resumes"
JOB_DIR = BASE_DIR / "jobs"

RESUME_DIR.mkdir(parents=True, exist_ok=True)
JOB_DIR.mkdir(parents=True, exist_ok=True)

# Don't make the small local LLM process too many jobs.
MAX_LLM_JOBS = 12

# Search results per query.
RESULTS_PER_QUERY = 10

# Maximum page text sent to Qwen.
MAX_LLM_PAGE_TEXT = 7000

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 "
        "(KHTML, like Gecko) "
        "Chrome/131.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "en-IN,en;q=0.9",
}


# ============================================================
# LOAD PROFILE
# ============================================================

if not PROFILE_FILE.exists():
    raise FileNotFoundError(
        f"\nprofile.json was not found.\n"
        f"Create it here:\n{PROFILE_FILE}\n"
    )

with open(PROFILE_FILE, "r", encoding="utf-8") as f:
    PROFILE = json.load(f)


# ============================================================
# SMALL HELPERS
# ============================================================

def clean_text(text):
    """Normalize whitespace."""
    if not text:
        return ""

    return re.sub(
        r"\s+",
        " ",
        text
    ).strip()


def normalize_url(url):
    """Remove tracking/query parameters where possible."""
    if not url:
        return ""

    parsed = urlparse(url)

    return (
        parsed.scheme
        + "://"
        + parsed.netloc
        + parsed.path
    ).rstrip("/")


# ============================================================
# WEB SEARCH
# PYTHON DOES THIS
# ============================================================

def web_search(query, max_results=RESULTS_PER_QUERY):

    print(f"\n🔎 Searching: {query}")

    results = []

    try:

        with DDGS() as ddgs:

            data = ddgs.text(
                query,
                region="in-en",
                safesearch="moderate",
                max_results=max_results
            )

            for item in data:

                title = clean_text(
                    item.get("title", "")
                )

                url = item.get(
                    "href",
                    ""
                )

                snippet = clean_text(
                    item.get("body", "")
                )

                if not url:
                    continue

                results.append({
                    "title": title,
                    "url": url,
                    "snippet": snippet,
                    "page_text": "",
                    "page_accessible": False
                })

    except Exception as e:

        print(
            f"❌ Search error: {e}"
        )

    return results


# ============================================================
# FETCH WEB PAGE
# PYTHON DOES THIS
#
# 403 is NOT fatal.
# Search result information is still retained.
# ============================================================

def fetch_page(url):

    try:

        response = requests.get(
            url,
            headers=HEADERS,
            timeout=12
        )

        if response.status_code == 403:

            print(
                f"⚠️ 403 blocked: {url}"
            )

            return ""

        if response.status_code >= 400:

            print(
                f"⚠️ HTTP {response.status_code}: {url}"
            )

            return ""

        soup = BeautifulSoup(
            response.text,
            "html.parser"
        )

        # Remove useless page elements.
        for tag in soup([
            "script",
            "style",
            "noscript",
            "svg",
            "nav",
            "footer"
        ]):
            tag.decompose()

        text = soup.get_text(
            separator=" ",
            strip=True
        )

        return clean_text(text)[:15000]

    except requests.RequestException as e:

        print(
            f"⚠️ Could not fetch page: {url}"
        )

        return ""


# ============================================================
# DETECT SEARCH/LISTING PAGES
#
# We DON'T want to send:
#
# "50 Fresher Hospital Jobs..."
#
# to Qwen as though it were an actual vacancy.
# ============================================================

def looks_like_listing_page(job):

    title = job["title"].lower()
    url = job["url"].lower()
    snippet = job["snippet"].lower()

    listing_title_patterns = [

        "jobs in",
        "jobs near",
        "job vacancies in",
        "vacancies in",
        "job openings in",
        "search jobs",
        "find jobs",
        "all jobs",
        "latest jobs",
        "fresher jobs in",
        "50 fresher",
        "100 fresher",
        "job vacancy",
        "job vacancies",
        "jobs and vacancies",
        "careers and jobs"
    ]

    for pattern in listing_title_patterns:

        if pattern in title:
            return True

    # Common search-page URL patterns.
    listing_url_patterns = [

        "/search?",
        "/search/",
        "/q-",
        "?q=",
        "/jobs-in-",
        "/jobs?",
        "/search-jobs",
        "/job-search"
    ]

    for pattern in listing_url_patterns:

        if pattern in url:
            return True

    # If title itself contains an obvious result-count phrase.
    if re.search(
        r"\b\d+\s+(?:fresher\s+)?(?:hospital\s+)?jobs\b",
        title
    ):
        return True

    return False


# ============================================================
# BASIC PYTHON JOB FILTER
#
# This happens BEFORE the LLM.
# ============================================================

def basic_filter(job, user_request):

    combined = clean_text(
        " ".join([
            job.get("title", ""),
            job.get("snippet", ""),
            job.get("page_text", "")
        ])
    ).lower()

    # --------------------------------------------------------
    # Reject obvious listing/search pages
    # --------------------------------------------------------

    if looks_like_listing_page(job):

        return False, "listing/search page"

    # --------------------------------------------------------
    # Reject obviously senior roles
    # --------------------------------------------------------

    senior_patterns = [

        "senior manager",
        "senior director",
        "general manager",
        "assistant general manager",
        "deputy general manager",
        "vice president",
        "chief executive",
        "10+ years",
        "8+ years",
        "7+ years",
        "6+ years"
    ]

    for pattern in senior_patterns:

        if pattern in combined:

            return False, f"senior requirement: {pattern}"

    # --------------------------------------------------------
    # Reject obvious unrelated jobs
    # --------------------------------------------------------

    # These are deliberately conservative.
    # We don't want Python rejecting a potentially useful job.
    obvious_bad_roles = [

        "sales manager",
        "relationship manager",
        "real estate agent",
        "telecaller",
        "insurance advisor"
    ]

    title = job["title"].lower()

    for role in obvious_bad_roles:

        if role in title:

            return False, f"unrelated role: {role}"

    return True, "passed Python filter"


# ============================================================
# DEDUPLICATION
# ============================================================

def deduplicate(jobs):

    seen = set()
    unique = []

    for job in jobs:

        normalized = normalize_url(
            job["url"]
        )

        if not normalized:
            continue

        key = hashlib.md5(
            normalized.encode(
                "utf-8"
            )
        ).hexdigest()

        if key in seen:
            continue

        seen.add(key)

        job["normalized_url"] = normalized

        unique.append(job)

    return unique


# ============================================================
# SAVE SEARCH RESULTS
# ============================================================

def save_jobs(jobs):

    output = JOB_DIR / "latest_search.json"

    with open(
        output,
        "w",
        encoding="utf-8"
    ) as f:

        json.dump(
            jobs,
            f,
            indent=2,
            ensure_ascii=False
        )

    print(
        f"\n💾 Saved search data: {output}"
    )


# ============================================================
# LLM JSON CALL
# ============================================================

def ask_llm_json(prompt):

    try:

        response = chat(
            model=MODEL,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            format="json"
        )

        content = response["message"]["content"]

        return json.loads(content)

    except json.JSONDecodeError as e:

        print(
            f"❌ LLM returned invalid JSON: {e}"
        )

        return None

    except Exception as e:

        print(
            f"❌ LLM error: {e}"
        )

        return None


# ============================================================
# LLM JOB EVALUATION
#
# Qwen only receives useful information.
# ============================================================

def evaluate_job(job, user_request):

    page_text = job.get(
        "page_text",
        ""
    )

    # If site blocked us, use search snippet.
    if page_text:

        evidence = page_text[
            :MAX_LLM_PAGE_TEXT
        ]

    else:

        evidence = job.get(
            "snippet",
            ""
        )[:3000]

    prompt = f"""
You are a careful job-matching assistant.

USER JOB REQUEST:
{user_request}

CANDIDATE PROFILE:
{json.dumps(PROFILE, ensure_ascii=False)}

VACANCY:
Title: {job['title']}
URL: {job['url']}

AVAILABLE JOB INFORMATION:
{evidence}

Determine whether this is a genuine and suitable vacancy.

IMPORTANT:
- Do NOT invent missing information.
- Do NOT assume a qualification the candidate does not have.
- Do NOT assume a vacancy is genuine if the information is insufficient.
- A fresher may match jobs explicitly allowing freshers.
- Consider qualification, experience, role, location and candidate experience.
- If this is actually a generic search/listing page, reject it.
- Score 0-100.

Return ONLY JSON:

{{
    "match": true,
    "score": 0,
    "confidence": 0,
    "reason": "short explanation",
    "missing_requirements": [],
    "resume_focus": []
}}
"""

    result = ask_llm_json(
        prompt
    )

    if not result:

        return {
            "match": False,
            "score": 0,
            "confidence": 0,
            "reason": "LLM evaluation failed",
            "missing_requirements": [],
            "resume_focus": []
        }

    return result


# ============================================================
# LLM RESUME TAILORING
#
# Only performed AFTER user chooses a vacancy.
# ============================================================

def tailor_resume(job):

    page_text = job.get(
        "page_text",
        ""
    )

    if not page_text:

        page_text = job.get(
            "snippet",
            ""
        )

    prompt = f"""
Create truthful resume content for this vacancy.

CANDIDATE PROFILE:
{json.dumps(PROFILE, ensure_ascii=False)}

JOB TITLE:
{job['title']}

JOB INFORMATION:
{page_text[:10000]}

RULES:
- NEVER invent qualifications.
- NEVER invent employers.
- NEVER invent job experience.
- NEVER invent certificates.
- NEVER invent dates.
- Use only facts present in the candidate profile.
- You may improve wording.
- You may reorder genuine skills.
- Keep it concise and professional.

Return ONLY JSON:

{{
    "summary": "...",
    "skills": [
        "...",
        "..."
    ],
    "experience_bullets": [
        "...",
        "..."
    ],
    "education": "..."
}}
"""

    return ask_llm_json(
        prompt
    )


# ============================================================
# CREATE RESUME
#
# Python creates the actual document.
# ============================================================

def create_resume(job, tailored):

    safe_name = re.sub(
        r"[^a-zA-Z0-9_-]+",
        "_",
        job["title"]
    ).strip("_")

    if not safe_name:

        safe_name = "tailored_resume"

    filename = (
        RESUME_DIR
        / f"{safe_name[:60]}.docx"
    )

    doc = Document()

    # --------------------------------------------------------
    # NAME
    # --------------------------------------------------------

    doc.add_heading(
        PROFILE.get(
            "name",
            "Candidate"
        ),
        level=0
    )

    contact = " | ".join(
        filter(
            None,
            [
                PROFILE.get(
                    "location",
                    ""
                ),
                PROFILE.get(
                    "phone",
                    ""
                ),
                PROFILE.get(
                    "email",
                    ""
                )
            ]
        )
    )

    if contact:

        doc.add_paragraph(
            contact
        )

    # --------------------------------------------------------
    # SUMMARY
    # --------------------------------------------------------

    doc.add_heading(
        "Professional Summary",
        level=1
    )

    doc.add_paragraph(
        tailored.get(
            "summary",
            ""
        )
    )

    # --------------------------------------------------------
    # SKILLS
    # --------------------------------------------------------

    skills = tailored.get(
        "skills",
        []
    )

    if skills:

        doc.add_heading(
            "Skills",
            level=1
        )

        for skill in skills:

            doc.add_paragraph(
                str(skill),
                style="List Bullet"
            )

    # --------------------------------------------------------
    # EXPERIENCE
    # --------------------------------------------------------

    doc.add_heading(
        "Experience",
        level=1
    )

    for experience in PROFILE.get(
        "experience",
        []
    ):

        role = experience.get(
            "role",
            ""
        )

        company = experience.get(
            "company",
            ""
        )

        location = experience.get(
            "location",
            ""
        )

        heading = " — ".join(
            filter(
                None,
                [
                    role,
                    company,
                    location
                ]
            )
        )

        doc.add_paragraph(
            heading
        )

    for bullet in tailored.get(
        "experience_bullets",
        []
    ):

        doc.add_paragraph(
            str(bullet),
            style="List Bullet"
        )

    # --------------------------------------------------------
    # EDUCATION
    # --------------------------------------------------------

    doc.add_heading(
        "Education",
        level=1
    )

    education = tailored.get(
        "education",
        ""
    )

    if education:

        doc.add_paragraph(
            education
        )

    doc.save(
        filename
    )

    print(
        f"\n📄 Resume created:\n{filename}"
    )

    return filename


# ============================================================
# SEARCH PIPELINE
# ============================================================

def find_jobs(user_request):

    # --------------------------------------------------------
    # Python generates several focused searches.
    # --------------------------------------------------------

    queries = [
        user_request,
        f"{user_request} vacancy",
        f"{user_request} jobs",
        f"{user_request} careers",
        f"{user_request} recruitment"
    ]

    all_jobs = []

    # --------------------------------------------------------
    # SEARCH
    # --------------------------------------------------------

    for query in queries:

        results = web_search(
            query,
            RESULTS_PER_QUERY
        )

        for result in results:

            if not result["url"]:
                continue

            # ------------------------------------------------
            # Reject obvious listing pages BEFORE fetching.
            # ------------------------------------------------

            if looks_like_listing_page(
                result
            ):

                print(
                    "↪ Skipping listing page:",
                    result["title"]
                )

                continue

            # ------------------------------------------------
            # Try fetching page.
            # 403 is okay.
            # ------------------------------------------------

            print(
                f"📄 Checking: "
                f"{result['title']}"
            )

            page_text = fetch_page(
                result["url"]
            )

            result["page_text"] = page_text

            result["page_accessible"] = bool(
                page_text
            )

            # ------------------------------------------------
            # Python filter.
            # ------------------------------------------------

            passed, reason = basic_filter(
                result,
                user_request
            )

            if passed:

                all_jobs.append(
                    result
                )

            else:

                print(
                    f"↪ Rejected: {reason}"
                )

    # --------------------------------------------------------
    # Remove duplicates.
    # --------------------------------------------------------

    all_jobs = deduplicate(
        all_jobs
    )

    # --------------------------------------------------------
    # Sort accessible pages first.
    # These contain more information.
    # --------------------------------------------------------

    all_jobs.sort(
        key=lambda x:
        x.get(
            "page_accessible",
            False
        ),
        reverse=True
    )

    save_jobs(
        all_jobs
    )

    print(
        "\n" + "=" * 60
    )

    print(
        "📊 Candidate vacancies after "
        "Python filtering:",
        len(all_jobs)
    )

    print(
        "=" * 60
    )

    return all_jobs


# ============================================================
# BROWSER APPLICATION
#
# Opens page, attempts resume upload,
# NEVER submits automatically.
# ============================================================

def open_application(
    url,
    resume_path
):

    print(
        "\n🌐 Opening application:"
    )

    print(
        url
    )

    with sync_playwright() as p:

        browser = p.chromium.launch(
            headless=False
        )

        page = browser.new_page()

        try:

            page.goto(
                url,
                wait_until="domcontentloaded",
                timeout=60000
            )

        except Exception as e:

            print(
                f"⚠️ Page navigation issue: {e}"
            )

        print(
            "\n🖥️ Browser opened."
        